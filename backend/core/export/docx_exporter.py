"""
backend/core/export/docx_exporter.py
──────────────────────────────────────
Exports a generated report dict to a professionally formatted DOCX.
Uses python-docx — 100% free, no external services.

Output: A Word document with:
  - Cover section (title, query, metadata)
  - All report sections with content
  - Sources reference list
"""
from __future__ import annotations
import io
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from backend.utils.logger import logger


# ── Color palette ─────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x3A, 0x6B)
BLUE   = RGBColor(0x25, 0x63, 0xEB)
TEAL   = RGBColor(0x0F, 0x76, 0x6E)
GRAY   = RGBColor(0x6B, 0x72, 0x80)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x11, 0x18, 0x27)


def export_report_to_docx(report: dict[str, Any], output_path: str | None = None) -> bytes:
    """
    Convert a report dict (from agent.generate_report) to DOCX bytes.

    Args:
        report: Report dict with title, query, sections, all_sources, etc.
        output_path: If provided, also saves to disk.

    Returns:
        DOCX file as bytes (for API streaming).
    """
    doc = Document()
    _set_page_margins(doc)
    _set_default_font(doc)

    # ── Cover ─────────────────────────────────────────────────────────────────
    _add_cover(doc, report)

    # ── Sections ──────────────────────────────────────────────────────────────
    for section in report.get("sections", []):
        _add_section(doc, section)

    # ── Sources ───────────────────────────────────────────────────────────────
    _add_sources(doc, report.get("all_sources", []))

    # ── Footer ────────────────────────────────────────────────────────────────
    _add_footer(doc, report.get("title", "Report"))

    # ── Serialise ─────────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    if output_path:
        Path(output_path).write_bytes(docx_bytes)
        logger.info(f"Report saved to {output_path}")

    logger.info(f"DOCX export complete | {len(docx_bytes):,} bytes")
    return docx_bytes


# ── Helpers ───────────────────────────────────────────────────────────────────
def _set_page_margins(doc: Document):
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)


def _set_default_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK


def _add_cover(doc: Document, report: dict):
    """Add a clean cover block."""
    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(report.get("title", "Executive Report"))
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = NAVY
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.space_after = Pt(6)

    # Subtitle line
    doc.add_paragraph()   # spacer
    sub = doc.add_paragraph()
    sub.add_run("AI-Generated Executive Report  ·  RAG Executive Analyst").font.color.rgb = GRAY
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata table
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.style = "Table Grid"
    cells = [
        ("Research Query", report.get("query", "")),
        ("Generated At", datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")),
        ("Sections", str(len(report.get("sections", [])))),
    ]
    for i, (label, value) in enumerate(cells):
        lc = meta_table.rows[i].cells[0]
        vc = meta_table.rows[i].cells[1]
        lc.paragraphs[0].add_run(label).bold = True
        lc.paragraphs[0].runs[0].font.color.rgb = NAVY
        vc.paragraphs[0].add_run(value)

    doc.add_paragraph()   # spacer
    doc.add_paragraph().add_run("─" * 80).font.color.rgb = GRAY
    doc.add_paragraph()   # spacer


def _add_section(doc: Document, section: dict):
    """Add a report section: heading + content."""
    # Section heading
    heading = doc.add_heading(section.get("section", "Section"), level=1)
    for run in heading.runs:
        run.font.color.rgb = NAVY

    # Content paragraphs
    content = section.get("content", "")
    for para_text in content.split("\n"):
        para_text = para_text.strip()
        if not para_text:
            continue

        # Numbered list items
        if para_text and para_text[0].isdigit() and ". " in para_text[:4]:
            p = doc.add_paragraph(style="List Number")
            p.add_run(para_text.split(". ", 1)[-1])

        # Bullet items
        elif para_text.startswith("- ") or para_text.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(para_text[2:])

        # Bold headers (lines ending with :)
        elif para_text.endswith(":") and len(para_text) < 60:
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.bold = True
            run.font.color.rgb = TEAL

        # Markdown table rows — simple conversion
        elif para_text.startswith("|"):
            p = doc.add_paragraph()
            p.add_run(para_text.replace("|", "   ")).font.name = "Courier New"
            p.runs[0].font.size = Pt(9)

        else:
            p = doc.add_paragraph()
            p.add_run(para_text)

    doc.add_paragraph()   # spacer after section


def _add_sources(doc: Document, sources: list[dict]):
    """Add a numbered sources reference section."""
    if not sources:
        return

    doc.add_page_break()
    heading = doc.add_heading("Sources & References", level=1)
    for run in heading.runs:
        run.font.color.rgb = NAVY

    # Deduplicate sources
    seen = set()
    unique_sources = []
    for s in sources:
        key = s.get("url") or f"{s.get('filename')}:{s.get('page')}"
        if key not in seen:
            seen.add(key)
            unique_sources.append(s)

    for i, source in enumerate(unique_sources, 1):
        p = doc.add_paragraph()
        num_run = p.add_run(f"[SOURCE {i}]  ")
        num_run.bold = True
        num_run.font.color.rgb = BLUE

        source_type = source.get("source_type", "internal")

        if source_type == "web":
            detail = f"{source.get('filename', 'Web Source')}  |  {source.get('url', '')}"
        else:
            detail = (
                f"{source.get('filename', 'Document')}  |  "
                f"Page {source.get('page', '?')}  |  "
                f"Relevance score: {source.get('score', 0):.3f}"
            )

        detail_run = p.add_run(detail)
        detail_run.font.color.rgb = GRAY
        detail_run.font.size = Pt(10)


def _add_footer(doc: Document, title: str):
    """Add page footer with doc title and date."""
    from docx.oxml.ns import qn
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = footer_para.add_run(
        f"{title}  ·  RAG Executive Analyst  ·  "
        f"{datetime.utcnow().strftime('%Y-%m-%d')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
