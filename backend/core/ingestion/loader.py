"""
backend/core/ingestion/loader.py
─────────────────────────────────
Loads PDF, DOCX, XLSX, and TXT files into raw text + metadata.
Returns a list of page-level dicts ready for chunking.
"""
from __future__ import annotations
import io
from pathlib import Path
from typing import Any

from backend.utils.logger import logger


def load_document(file_path: str | Path, filename: str) -> list[dict[str, Any]]:
    """
    Load any supported document type.

    Returns:
        List of page dicts: {"text": str, "page": int, "metadata": dict}
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    loaders = {
        ".pdf":  _load_pdf,
        ".docx": _load_docx,
        ".doc":  _load_docx,
        ".xlsx": _load_xlsx,
        ".xls":  _load_xlsx,
        ".txt":  _load_txt,
        ".md":   _load_txt,
    }

    loader_fn = loaders.get(suffix)
    if not loader_fn:
        raise ValueError(f"Unsupported file type: {suffix}. Supported: {list(loaders)}")

    logger.info(f"Loading {suffix} file: {filename}")
    pages = loader_fn(path, filename)
    logger.info(f"  → Loaded {len(pages)} page(s) from {filename}")
    return pages


# ── PDF ───────────────────────────────────────────────────────────────────────
def _load_pdf(path: Path, filename: str) -> list[dict]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()
        if len(text) < 20:          # Skip near-empty pages
            continue
        pages.append({
            "text": text,
            "page": i + 1,
            "metadata": {
                "filename": filename,
                "source_type": "internal",
                "file_type": "pdf",
                "page": i + 1,
                "total_pages": len(reader.pages),
            }
        })
    return pages


# ── DOCX ──────────────────────────────────────────────────────────────────────
def _load_docx(path: Path, filename: str) -> list[dict]:
    from docx import Document

    doc = Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # Split into ~page-sized chunks (every 3000 chars ≈ 1 page)
    page_size = 3000
    pages = []
    for i, start in enumerate(range(0, max(len(full_text), 1), page_size)):
        chunk_text = full_text[start:start + page_size].strip()
        if not chunk_text:
            continue
        pages.append({
            "text": chunk_text,
            "page": i + 1,
            "metadata": {
                "filename": filename,
                "source_type": "internal",
                "file_type": "docx",
                "page": i + 1,
                "total_pages": max(1, len(full_text) // page_size),
            }
        })

    # Also extract tables as text
    for tbl_i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        table_text = "\n".join(rows)
        if table_text.strip():
            pages.append({
                "text": f"[TABLE {tbl_i+1}]\n{table_text}",
                "page": tbl_i + 1,
                "metadata": {
                    "filename": filename,
                    "source_type": "internal",
                    "file_type": "docx_table",
                    "page": tbl_i + 1,
                    "total_pages": len(doc.tables),
                }
            })
    return pages


# ── XLSX ──────────────────────────────────────────────────────────────────────
def _load_xlsx(path: Path, filename: str) -> list[dict]:
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True)
    pages = []
    for sheet_i, sheet in enumerate(wb.worksheets):
        rows = []
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(c) if c is not None else "" for c in row)
            if row_text.strip().replace("|", "").strip():
                rows.append(row_text)
        sheet_text = f"[SHEET: {sheet.title}]\n" + "\n".join(rows)
        if rows:
            pages.append({
                "text": sheet_text,
                "page": sheet_i + 1,
                "metadata": {
                    "filename": filename,
                    "source_type": "internal",
                    "file_type": "xlsx",
                    "page": sheet_i + 1,
                    "total_pages": len(wb.worksheets),
                    "sheet": sheet.title,
                }
            })
    return pages


# ── TXT / MD ──────────────────────────────────────────────────────────────────
def _load_txt(path: Path, filename: str) -> list[dict]:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    page_size = 3000
    pages = []
    for i, start in enumerate(range(0, max(len(text), 1), page_size)):
        chunk_text = text[start:start + page_size].strip()
        if chunk_text:
            pages.append({
                "text": chunk_text,
                "page": i + 1,
                "metadata": {
                    "filename": filename,
                    "source_type": "internal",
                    "file_type": path.suffix.lstrip("."),
                    "page": i + 1,
                    "total_pages": max(1, len(text) // page_size),
                }
            })
    return pages
